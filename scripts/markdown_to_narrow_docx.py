#!/usr/bin/env python3
"""Convert a Markdown quick-reference manual to a narrow-margin DOCX."""

from __future__ import annotations

import argparse
import platform
import re
from pathlib import Path


def default_font_for_platform() -> str:
    system = platform.system().lower()
    if system == "windows":
        return "SimSun"
    if system == "darwin":
        return "Songti SC"
    return "Noto Serif CJK SC"


def require_docx():
    try:
        import docx  # type: ignore
        from docx import Document  # type: ignore
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.shared import Inches, Pt  # type: ignore
        from docx.enum.text import WD_LINE_SPACING  # type: ignore
    except Exception as exc:
        raise SystemExit("python-docx is required. Install it in the active environment or use the documents runtime.") from exc
    return docx, Document, OxmlElement, qn, Inches, Pt, WD_LINE_SPACING


def set_run_font(run, qn, name: str, size_pt: float | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size_pt is not None:
        from docx.shared import Pt  # type: ignore
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def add_runs_with_bold(paragraph, text: str, qn, font_name: str, size_pt: float):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        clean = part[2:-2] if is_bold else part
        run = paragraph.add_run(clean)
        set_run_font(run, qn, font_name, size_pt, is_bold)


def configure_document(document, qn, Inches, Pt, WD_LINE_SPACING, font_name: str, margin_inches: float):
    section = document.sections[0]
    section.top_margin = Inches(margin_inches)
    section.bottom_margin = Inches(margin_inches)
    section.left_margin = Inches(margin_inches)
    section.right_margin = Inches(margin_inches)

    styles = document.styles
    for style_name, size, bold in [
        ("Normal", 8.5, False),
        ("Heading 1", 15, True),
        ("Heading 2", 12, True),
        ("Heading 3", 10.5, True),
        ("List Bullet", 8.5, False),
    ]:
        style = styles[style_name]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(size)
        style.font.bold = bold
        if hasattr(style, "paragraph_format"):
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(2)
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_markdown_line(document, line: str, qn, font_name: str):
    if not line.strip():
        return
    stripped = line.rstrip()
    heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
    if heading:
        level = min(len(heading.group(1)), 3)
        paragraph = document.add_heading(level=level)
        add_runs_with_bold(paragraph, heading.group(2), qn, font_name, 15 if level == 1 else 12 if level == 2 else 10.5)
        return

    bullet = re.match(r"^(\s*)[-*+]\s+(.*)$", stripped)
    numbered = re.match(r"^(\s*)\d+[.)]\s+(.*)$", stripped)
    if bullet or numbered:
        match = bullet or numbered
        indent_spaces = len(match.group(1).replace("\t", "    "))
        text = match.group(2)
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = None
        if indent_spaces:
            from docx.shared import Inches  # type: ignore
            paragraph.paragraph_format.left_indent = Inches(min(0.5, indent_spaces * 0.035))
        add_runs_with_bold(paragraph, text, qn, font_name, 8.5)
        return

    paragraph = document.add_paragraph()
    add_runs_with_bold(paragraph, stripped, qn, font_name, 8.5)


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert Markdown quick-reference text to narrow-margin DOCX.")
    parser.add_argument("markdown", type=Path)
    parser.add_argument("docx", type=Path)
    parser.add_argument(
        "--font",
        default=default_font_for_platform(),
        help="Body font. Defaults by OS: Windows SimSun, macOS Songti SC, Linux Noto Serif CJK SC.",
    )
    parser.add_argument("--margin", type=float, default=0.35, help="Margins in inches")
    args = parser.parse_args()

    _, Document, _, qn, Inches, Pt, WD_LINE_SPACING = require_docx()
    document = Document()
    configure_document(document, qn, Inches, Pt, WD_LINE_SPACING, args.font, args.margin)
    text = args.markdown.expanduser().read_text(encoding="utf-8")
    for line in text.splitlines():
        add_markdown_line(document, line, qn, args.font)
    args.docx.expanduser().parent.mkdir(parents=True, exist_ok=True)
    document.save(str(args.docx.expanduser()))
    print(str(args.docx.expanduser().resolve()))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
